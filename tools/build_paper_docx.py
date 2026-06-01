from __future__ import annotations

from pathlib import Path
import re
from typing import Any, cast

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = ROOT / "paper"
SOURCE_MD = PAPER_DIR / "paper_draft.md"
TEMPLATE = PAPER_DIR / "paper_template_original.docx"
if not TEMPLATE.exists():
    raise FileNotFoundError(f"Template DOCX not found: {TEMPLATE}")
OUTPUT = PAPER_DIR / "paper_submission_STRIDE_ZAP.docx"
IMAGE_PATTERN = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")


def set_run_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Batang")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_paragraph(doc: DocxDocument, text: str = "", size: int = 11, bold: bool = False, align=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def clear_document_body(doc: DocxDocument) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_document(doc: DocxDocument) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    for style_name in ("Normal", "Body Text"):
        try:
            style = cast(Any, doc.styles[style_name])
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Batang")
        style.font.size = Pt(11)


def is_markdown_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace(":", "").replace("-", "").strip()) == set()


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    headers = [cell.strip() for cell in lines[start].strip().strip("|").split("|")]
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append([cell.strip().replace("`", "") for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return headers, rows, index


def set_cell_text(cell, text: str, size: int = 9, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: DocxDocument, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=9, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, size=9, align=align)

    add_paragraph(doc, "", size=9)


def add_image(doc: DocxDocument, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))
    paragraph.paragraph_format.space_after = Pt(3)


def emit_markdown(doc: DocxDocument, text: str) -> None:
    lines = text.splitlines()
    index = 0
    pending_table_captions: list[str] = []
    in_front_matter = True

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            add_paragraph(doc, line[2:], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith("## "):
            heading_text = line[3:]
            if heading_text == "초록":
                in_front_matter = False
            add_paragraph(doc, heading_text, size=11, bold=True)
        elif line.startswith("### "):
            add_paragraph(doc, line[4:], size=11, bold=True)
        elif line.startswith("[표 ") or line.startswith("[Table "):
            pending_table_captions.append(line)
        elif line.startswith("[그림 ") or line.startswith("[Figure "):
            add_paragraph(doc, line, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith("!["):
            image_match = IMAGE_PATTERN.fullmatch(line)
            if image_match:
                image_path = (SOURCE_MD.parent / image_match.group("path")).resolve()
                if image_path.exists():
                    add_image(doc, image_path)
                else:
                    add_paragraph(doc, f"[이미지 누락: {image_match.group('path')}]", size=9)
        elif line.startswith("|") and index + 1 < len(lines) and is_markdown_table_separator(lines[index + 1]):
            for caption in pending_table_captions:
                add_paragraph(doc, caption, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            pending_table_captions.clear()
            headers, rows, index = parse_table(lines, index)
            add_table(doc, headers, rows)
            continue
        else:
            if pending_table_captions:
                for caption in pending_table_captions:
                    add_paragraph(doc, caption, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                pending_table_captions.clear()
            size = 10 if line.startswith("[") and "]" in line[:5] else 11
            align = WD_ALIGN_PARAGRAPH.CENTER if in_front_matter else None
            front_size = 10 if in_front_matter else size
            add_paragraph(doc, line.replace("`", ""), size=front_size, align=align)

        index += 1


def build() -> None:
    doc = create_document(str(TEMPLATE))
    clear_document_body(doc)
    configure_document(doc)
    emit_markdown(doc, SOURCE_MD.read_text(encoding="utf-8"))
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build()
